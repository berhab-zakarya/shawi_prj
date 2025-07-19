from docx import Document
from django.template.loader import render_to_string

def generate_docx(contract):
    doc = Document()
    doc.add_heading(f'{contract.contract_type} Contract', 0)
    doc.add_paragraph(f'Client: {contract.data.get("client_name", "")}')
    doc.add_paragraph(f'Details: {contract.data.get("details", "")}')
    doc.add_paragraph(f'Generated on: {contract.created_at}')
    return doc